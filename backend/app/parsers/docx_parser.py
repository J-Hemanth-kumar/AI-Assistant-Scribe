from __future__ import annotations

import io
from typing import Any

from docx import Document as DocxDocument

from app.parsers.base import ParsedBlock


def parse_docx(data: bytes) -> list[ParsedBlock]:
    doc = DocxDocument(io.BytesIO(data))
    blocks: list[ParsedBlock] = []

    block_index = 0
    current_section: str | None = None

    def _extract_run_font_size(paragraph_text) -> tuple[str | None, float | None]:
        # python-docx doesn't expose consistent font metadata at paragraph level;
        # we do a best-effort extraction from the first run that has font info.
        for run in getattr(paragraph_text, "runs", []) or []:
            if run is None or run.font is None:
                continue
            font_name = run.font.name if run.font.name else None
            size_pt = float(run.font.size.pt) if run.font.size is not None else None
            if font_name or size_pt:
                return font_name, size_pt
        return None, None

    def _heading_prefix(style_name: str) -> str | None:
        s = (style_name or "").strip().lower()
        # Common Word styles: "Heading 1", "Heading 2", ...
        if s.startswith("heading"):
            parts = s.split()
            # Extract first number token if present.
            for p in parts:
                if p.isdigit():
                    level = int(p)
                    return "#" * max(1, min(level, 6)) + " "
        return None

    # Paragraphs (including headings)
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if not text:
            continue

        bbox = None  # DOCX doesn't provide bounding boxes without extra tooling.

        style_name = paragraph.style.name if getattr(paragraph, "style", None) is not None else ""
        heading_prefix = _heading_prefix(style_name)
        if heading_prefix:
            # Preserve heading hierarchy in text, and also expose `section` for downstream RAG.
            text = f"{heading_prefix}{text}"
            current_section = (paragraph.text or "").strip()

        font_name, size_pt = _extract_run_font_size(paragraph)

        blocks.append(
            ParsedBlock(
                page_no=None,
                block_index=block_index,
                text=text,
                parser_type="docx",
                section=current_section,
                bbox=bbox,
                font=font_name,
                size=size_pt,
            )
        )
        block_index += 1

    # Tables (flatten rows to pipe-separated blocks)
    for table in doc.tables:
        for row in table.rows:
            cell_texts: list[str] = []
            for cell in row.cells:
                # Flatten cell paragraphs; omit empty lines.
                parts = [(p.text or "").strip() for p in cell.paragraphs]
                parts = [p for p in parts if p]
                cell_texts.append(" ".join(parts))

            row_text = " | ".join([t for t in cell_texts if t])
            row_text = row_text.strip()
            if not row_text:
                continue

            # Best-effort: infer font/size from first paragraph in first non-empty cell.
            font_name = None
            size_pt = None
            for cell in row.cells:
                if cell.paragraphs:
                    font_name, size_pt = _extract_run_font_size(cell.paragraphs[0])
                    break

            blocks.append(
                ParsedBlock(
                    page_no=None,
                    block_index=block_index,
                    text=row_text,
                    parser_type="docx",
                    section=current_section,
                    bbox=None,
                    font=font_name,
                    size=size_pt,
                )
            )
            block_index += 1

    return blocks

