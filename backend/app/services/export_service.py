"""
Export Service — applies structured edits and exports to DOCX, PDF, MD, TXT.

Two types of edits are supported:

  Content edit  (updated_text != "")
    The text of the chunk is replaced with updated_text.
    Applied first so styling is layered on top of correct content.

  Style edit  (edit["style"] is not None)
    Font name, size, bold, italic, underline, alignment, line spacing.
    Applied via python-docx run/paragraph formatting (DOCX) or PyMuPDF
    font insertion (PDF).  Style metadata is NEVER stored as HTML — it
    lives in the structured "style" dict from EditStyle schema.

Export paths:
  DOCX (preferred for styled output):
    Reassembles full_text paragraph-by-paragraph, then applies style dicts.
    python-docx gives pixel-perfect font/size/bold control.

  PDF:
    Layout-preserving redact-and-replace via PyMuPDF.
    Font styling applied via fontname parameter on insert_text.
    Fallback: generate fresh PDF from full_text when redaction fails.

  MD / TXT:
    Plain-text export of full_text — styles are not representable.
"""
from __future__ import annotations

import logging
import re
from datetime import datetime, timezone
from io import BytesIO
from typing import Any

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

from app.services.minio_service import MinioService
from app.db.session import SessionLocal

logger = logging.getLogger(__name__)

# Map alignment strings to python-docx enum values
_ALIGN_MAP = {
    "left":    WD_ALIGN_PARAGRAPH.LEFT,
    "center":  WD_ALIGN_PARAGRAPH.CENTER,
    "right":   WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _apply_style_to_run(run, style: dict) -> None:
    """Apply EditStyle dict to a python-docx Run."""
    if style.get("font_name"):
        run.font.name = style["font_name"]
    if style.get("font_size") is not None:
        run.font.size = Pt(float(style["font_size"]))
    if style.get("bold") is not None:
        run.font.bold = style["bold"]
    if style.get("italic") is not None:
        run.font.italic = style["italic"]
    if style.get("underline") is not None:
        run.font.underline = style["underline"]
    if style.get("color_hex"):
        hex_val = style["color_hex"].lstrip("#")
        if len(hex_val) == 6:
            r, g, b = int(hex_val[0:2], 16), int(hex_val[2:4], 16), int(hex_val[4:6], 16)
            run.font.color.rgb = RGBColor(r, g, b)


def _apply_style_to_paragraph(para, style: dict) -> None:
    """Apply alignment and line spacing to a python-docx Paragraph."""
    if style.get("alignment") and style["alignment"] in _ALIGN_MAP:
        para.alignment = _ALIGN_MAP[style["alignment"]]
    if style.get("line_spacing") is not None:
        para.paragraph_format.line_spacing = float(style["line_spacing"])


def _build_style_index(edits: list[dict]) -> dict[int, dict]:
    """
    Build a lookup table  chunk_index -> style dict  from the edits list.
    Only entries with a non-None style are included.
    """
    return {
        e["chunk_index"]: e["style"]
        for e in edits
        if e.get("style") is not None
    }


def _build_content_index(edits: list[dict]) -> dict[int, str]:
    """
    Build chunk_index -> updated_text  for content edits only.
    Style-only edits (updated_text == "") are excluded.
    """
    return {
        e["chunk_index"]: e["updated_text"]
        for e in edits
        if e.get("updated_text", "").strip()
    }


# ---------------------------------------------------------------------------
# PDF helpers (unchanged from original, kept for layout-preserving path)
# ---------------------------------------------------------------------------

def _robust_token_replace(content: str, search_text: str, replacement_text: str) -> str:
    if not search_text:
        return content
    try:
        exact_re = re.compile(re.escape(search_text), re.IGNORECASE)
        if exact_re.search(content):
            return exact_re.sub(replacement_text, content, count=1)
    except Exception:
        pass
    search_tokens = re.findall(r"\w+", search_text)
    if not search_tokens:
        return content.replace(search_text, replacement_text)
    pos, match_start, match_end = 0, -1, -1
    for i, token in enumerate(search_tokens):
        m = re.search(re.escape(token), content[pos:], re.IGNORECASE)
        if not m:
            return content.replace(search_text, replacement_text)
        if i == 0:
            match_start = pos + m.start()
        pos += m.end()
        match_end = pos
    if match_start != -1 and match_end != -1:
        return content[:match_start] + replacement_text + content[match_end:]
    return content.replace(search_text, replacement_text)


# ---------------------------------------------------------------------------
# ExportService
# ---------------------------------------------------------------------------

class ExportService:
    def __init__(self) -> None:
        self.minio = MinioService()

    # ------------------------------------------------------------------
    # DOCX export  (preferred for styled output)
    # ------------------------------------------------------------------

    def _export_docx(
        self,
        full_text: str,
        edits: list[dict],
        doc_id: str,
        version_id: int,
    ) -> bytes:
        """
        Build a fully styled DOCX from full_text + structured edit metadata.

        Strategy:
          1. Split full_text on blank lines into paragraphs.
          2. Match each paragraph against the edit list by text proximity.
          3. Apply content edits (updated_text) and style edits (style dict).
          4. Headings (is_heading=True) are rendered as Heading 1 style.
        """
        doc = DocxDocument()

        # Pre-build lookup indexes
        style_idx   = _build_style_index(edits)
        content_idx = _build_content_index(edits)

        # Build a text→chunk_index reverse map for matching paragraphs
        chunk_text_map: dict[str, int] = {}
        for e in edits:
            orig = e.get("original_text", "").strip()
            if orig:
                chunk_text_map[orig] = e["chunk_index"]

        paragraphs = [p.strip() for p in full_text.split("\n\n") if p.strip()]

        for para_text in paragraphs:
            # Find the best matching edit for this paragraph
            matched_chunk: int | None = None
            for orig_text, chunk_idx in chunk_text_map.items():
                # Substring match — paragraph may contain just part of the chunk
                if orig_text in para_text or para_text in orig_text:
                    matched_chunk = chunk_idx
                    break

            # Determine final text and style
            final_text = content_idx.get(matched_chunk, para_text) if matched_chunk is not None else para_text
            style_dict = style_idx.get(matched_chunk, {}) if matched_chunk is not None else {}

            is_heading = style_dict.get("is_heading", False)

            if is_heading:
                para = doc.add_heading(final_text, level=1)
                # Apply style to the heading run so color/font takes effect
                if style_dict and para.runs:
                    _apply_style_to_run(para.runs[0], style_dict)
            else:
                para = doc.add_paragraph()
                run = para.add_run(final_text)
                if style_dict:
                    _apply_style_to_run(run, style_dict)

            if style_dict:
                _apply_style_to_paragraph(para, style_dict)

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ------------------------------------------------------------------
    # PDF helpers
    # ------------------------------------------------------------------

    def apply_edits_to_pdf(self, pdf_bytes: bytes, edits: list[dict]) -> tuple[bytes, bool]:
        """
        Layout-preserving redact-and-replace for PDFs.
        Now respects style.font_name when inserting replacement text.
        """
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        any_miss = False

        for edit in edits:
            raw_page_no = edit.get("page_no", 1)
            page_no = max(0, raw_page_no - 1) if raw_page_no > 0 else 0
            original_text = edit.get("original_text", "").strip()
            # For PDF: use updated_text if present, else keep original (style-only)
            updated_text = edit.get("updated_text", "").strip() or original_text
            style = edit.get("style") or {}

            if not original_text or page_no >= len(doc):
                continue

            page = doc[page_no]
            instances = page.search_for(original_text, quads=True)

            if not instances:
                page_words = page.get_text("words")
                search_words = re.findall(r"\w+", original_text)
                if search_words:
                    for i in range(len(page_words) - len(search_words) + 1):
                        if all(
                            re.sub(r"\W+", "", page_words[i + j][4]).lower() == search_words[j].lower()
                            for j in range(len(search_words))
                        ):
                            rects = [fitz.Rect(page_words[i + k][:4]) for k in range(len(search_words))]
                            combined = rects[0]
                            for r in rects[1:]:
                                combined.include_rect(r)
                            instances.append(combined.quad)
                            break

            if not instances:
                logger.warning("PDF: text not found on page %d: '%s'", page_no + 1, original_text[:50])
                any_miss = True
                continue

            # Resolve font for PyMuPDF
            # PyMuPDF built-in font names are limited; map common names
            fontname = "helv"  # default Helvetica
            if style.get("font_name"):
                fn = style["font_name"].lower().replace(" ", "")
                if "times" in fn or "roman" in fn:
                    fontname = "Times-Roman"
                elif "courier" in fn or "mono" in fn:
                    fontname = "Courier"
                # Others fall back to helv

            for quad in instances:
                page.add_redact_annot(quad, fill=(1, 1, 1))
            page.apply_redactions()

            # Always redraw: for content edits use updated_text,
            # for style-only edits use original_text (same content, new styling)
            redraw_text = (
                edit.get("updated_text", "").strip() or original_text
            )
            if redraw_text:
                first_rect = instances[0].rect
                font_size = style.get("font_size") or max(
                    abs(first_rect.y1 - first_rect.y0) * 0.85, 10
                )
                # Resolve text colour from style.color_hex (default black)
                text_color = (0.0, 0.0, 0.0)
                if style.get("color_hex"):
                    hex_val = style["color_hex"].lstrip("#")
                    if len(hex_val) == 6:
                        try:
                            text_color = (
                                int(hex_val[0:2], 16) / 255.0,
                                int(hex_val[2:4], 16) / 255.0,
                                int(hex_val[4:6], 16) / 255.0,
                            )
                        except ValueError:
                            pass

                page.insert_text(
                    (first_rect.x0, first_rect.y1 - 2),
                    redraw_text,
                    fontsize=float(font_size),
                    color=text_color,
                    fontname=fontname,
                )

        edited_bytes = doc.tobytes()
        doc.close()
        return edited_bytes, not any_miss

    def _generate_pdf_from_text(self, pages_data: list) -> bytes:
        """Plain-text fallback PDF generation (unchanged)."""
        doc = fitz.open()
        font_size = 11
        margin = 50
        for p_data in pages_data:
            text = p_data.get("text", "")
            page = doc.new_page()
            y_pos = margin
            lines: list[str] = []
            for paragraph in text.splitlines():
                words = paragraph.split()
                line = ""
                for w in words:
                    if (len(line) + len(w)) * 6 > (page.rect.width - 2 * margin):
                        lines.append(line)
                        line = w
                    else:
                        line = (line + " " + w).strip()
                lines.append(line)
            for line in lines:
                if y_pos > page.rect.height - margin:
                    page = doc.new_page()
                    y_pos = margin
                page.insert_text((margin, y_pos), line, fontsize=font_size, fontname="helv")
                y_pos += font_size * 1.3
        pdf_bytes = doc.tobytes()
        doc.close()
        return pdf_bytes

    # ------------------------------------------------------------------
    # Main export entry point
    # ------------------------------------------------------------------

    def export_with_edits(
        self,
        doc_id: str,
        version_id: int,
        original_key: str,
        edits: list[dict],
        export_format: str = "pdf",
    ) -> str:
        logger.info(
            "Export: doc_id=%s format=%s edits=%d", doc_id, export_format, len(edits)
        )

        original_bytes = self.minio.get_object_bytes(original_key)
        is_pdf = original_key.lower().endswith(".pdf")

        # Resolve full_text from stored version
        with SessionLocal() as db:
            from app.db.models import DocumentVersion
            dv = db.query(DocumentVersion).filter(DocumentVersion.id == version_id).first()
            full_text = dv.full_text or "" if dv else original_bytes.decode("utf-8", errors="ignore")

        # ------------------------------------------------------------------
        # DOCX — best format for styled output (font, size, bold all work)
        # ------------------------------------------------------------------
        if export_format == "docx":
            file_bytes = self._export_docx(full_text, edits, doc_id, version_id)
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ext = "docx"

        # ------------------------------------------------------------------
        # PDF — layout-preserving redact, fallback to fresh generation
        # ------------------------------------------------------------------
        elif export_format == "pdf":
            if is_pdf:
                file_bytes, success = self.apply_edits_to_pdf(original_bytes, edits)
                if not success:
                    logger.info("PDF redaction incomplete — falling back to text generation")
                    file_bytes = self._generate_pdf_from_text([{"page_no": 1, "text": full_text}])
            else:
                file_bytes = self._generate_pdf_from_text([{"page_no": 1, "text": full_text}])
            content_type = "application/pdf"
            ext = "pdf"

        # ------------------------------------------------------------------
        # Markdown
        # ------------------------------------------------------------------
        elif export_format in ("markdown", "md"):
            content = (
                f"# Exported Document\n\nDoc ID: {doc_id}\n"
                f"Version: {version_id}\n\n{full_text}"
            )
            file_bytes = content.encode("utf-8")
            content_type = "text/markdown"
            ext = "md"

        # ------------------------------------------------------------------
        # Plain text fallback
        # ------------------------------------------------------------------
        else:
            file_bytes = full_text.encode("utf-8")
            content_type = "text/plain"
            ext = "txt"

        key = f"exports/{doc_id}/{version_id}.{ext}"
        self.minio.upload_fileobj(BytesIO(file_bytes), key, content_type)
        return f"/api/v1/export/download/{key}"